import difflib
import io
from typing import Any, List

from docx import Document
from docx.shared import RGBColor

from reworkd_platform.schemas.workflow.base import Block, BlockIOBase
from reworkd_platform.services.aws.s3 import SimpleStorageService
from reworkd_platform.services.sockets import websockets
from reworkd_platform.settings import settings
from reworkd_platform.services.url_shortener import UrlShortenerService


class DiffDocInput(BlockIOBase):
    original: str
    updated: str


class DiffDocOutput(BlockIOBase):
    file_url: str


class DiffDoc(Block):
    type = "DiffDoc"
    description = (
        "Create a Word document that shows the difference between two bodies of text"
    )
    input: DiffDocInput

    async def run(self, workflow_id: str, **kwargs: Any) -> DiffDocOutput:
        with io.BytesIO() as diff_doc_file:
            websockets.log(workflow_id, "Creating diff of original and updated text")
            diffs = get_diff(self.input.original, self.input.updated)
            diff_doc_file = get_diff_doc(diffs, diff_doc_file)

            websockets.log(workflow_id, "Uploading diff doc to S3")
            s3_service = SimpleStorageService(settings.s3_bucket_name)
            s3_service.upload_to_bucket(
                object_name=f"docs/{workflow_id}/{self.id}.docx",
                file=diff_doc_file,
            )

            file_url = s3_service.create_presigned_download_url(
                object_name=f"docs/{workflow_id}/{self.id}.docx",
            )
            websockets.log(workflow_id, f"Diff Doc successfully uploaded to S3")

            shortener = UrlShortenerService()
            tiny_url = await shortener.get_shortened_url(file_url)
            websockets.log(workflow_id, f"Download the diff doc via: {tiny_url}")

            return DiffDocOutput(file_url=tiny_url)


def get_diff(original: str, updated: str) -> List[List[str]]:
    original_paragraphs = original.split("\n")
    updated_paragraphs = updated.split("\n")

    diffs = []
    for orig_par, updt_par in zip(original_paragraphs, updated_paragraphs):
        differ = difflib.Differ()
        words1 = orig_par.split()
        words2 = updt_par.split()

        diffs.append(list(differ.compare(words1, words2)))

    return diffs


def get_diff_doc(diff_list: List[List[str]], in_memory_file: io.BytesIO) -> io.BytesIO:
    """
    Create a Word document that shows the difference between two bodies of text.
    Each element of diff_list is a list of strings of type "  word", "- word", or "+ word".
    """

    doc = Document()

    for diff in diff_list:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = 0

        for word in diff:
            if word.startswith("  "):
                run = paragraph.add_run(word[2:] + " ")
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black color
            elif word.startswith("- "):
                run = paragraph.add_run(word[2:] + " ")
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red color
            elif word.startswith("+ "):
                run = paragraph.add_run(word[2:] + " ")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green color
            else:
                continue

    doc.save(in_memory_file)

    in_memory_file.seek(0)
    return in_memory_file
